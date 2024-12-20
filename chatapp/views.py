import json
import os
import re
from asgiref.sync import sync_to_async
from django.conf import settings
from django.http import FileResponse, HttpResponse, HttpResponseForbidden, JsonResponse
from django.contrib.auth.decorators import login_required
from django.shortcuts import render

from lorem_text import lorem
from box8.ChatAgent import chat, chat_enhance, chat_summarize, extract_insights, delete_entry,chat_save_conversation
from box8.utils_pdf import PdfUtils

import markdown


"""
Naldia.async_mode = False

chat = Chatter()
# l'objet commun ne sert qu'à vérifier la présence d'une vectorisation antérieure, voir new_async_memorizer
memorizer = Memorizer()

# deprecated : async_memorize = sync_to_async(memorizer.memorize)
async_load = sync_to_async(chat.load)
async_ask = sync_to_async(chat.ask)
"""


"""résumer la fiche référence avec une proposition de localisation géographique dans un json selon le format : {"résumé" : résumé, lat : latitude, lng : longitude}. Ne retourner que le json"""
"""résumer la mission avec une proposition de localisation géographique dans un json selon le format : {"résumé" : résumé, lat : latitude, lng : longitude}"""

# FUNCTIONS
def get_absolute_path(relative_path):
    parts = relative_path.split('/')
    file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), *parts)
    return file_path

# USER VIEWS
@login_required
def chatapp_dashboard(request):
    # initialisation des variables session 
    chatapp_init_session(request)
    context = {
        "welcome": "Assistant d'analyse documentaire"
    }
    return render(request, 'chatapp/dashboard.html', context=context)

@login_required
def chatapp_webscrapping_demo(request):
    # initialisation des variables session 
    chatapp_init_session(request)
    context = {
        "welcome": "Collecteur de données externes"
    }
    return render(request, 'chatapp/webscrapping.html', context=context)


@login_required
def models_dashboard(request):
    # initialisation des variables session 
    chatapp_init_session(request)
    context = {
        "welcome": "Liste des modèles"
    }
    return render(request, 'chatapp/models_dashboard.html', context=context)


@login_required
def models_dpgf_demo(request):
    # initialisation des variables session 
    chatapp_init_session(request)
    context = {
        "welcome": "Economie de la construction"
    }
    return render(request, 'chatapp/models_dpgf.html', context=context)

@login_required
def models_vision_demo(request):
    # initialisation des variables session 
    chatapp_init_session(request)
    context = {
        "welcome": "Vision par ordinateur"
    }
    return render(request, 'chatapp/models_vision.html', context=context)




@login_required
def models_dpgf_demo_upload(request):
    if request.method == 'POST' and request.FILES.getlist('files'):
        uploaded_files = request.FILES.getlist('files')
        
        success_count = 0
        failure_count = 0
        datasetPath = os.path.join(get_absolute_path("sharepoint"), "datasets")

        for uploaded_file in uploaded_files:
            destination_path = os.path.join(datasetPath, uploaded_file.name)
            try:
                with open(destination_path, 'wb+') as destination:
                    for chunk in uploaded_file.chunks():
                        destination.write(chunk)
                success_count += 1
            except Exception as e:
                failure_count += 1

        if success_count > 0:
            message = f'{success_count} fichiers téléchargés avec succès.'
        else:
            message = 'Échec du téléchargement des fichiers.'

        if failure_count > 0:
            message += f' {failure_count} fichiers n\'ont pas pu être téléchargés.'
            return JsonResponse({'message': 'Aucun fichier n\'a été téléchargé.'}, status=400)
        
        return JsonResponse({'message': 'Fichiers téléchargés.'})

    return JsonResponse({'message': 'Aucun fichier n\'a été téléchargé.'}, status=400)


@login_required
def models_dpgf_demo_train(request):
    data = json.loads(request.body.decode('utf-8'))
    modelName = data.get('model', '') # question posé
    
    datasetPath = os.path.join(get_absolute_path("sharepoint"), "datasets")
    response = "train_model(datasetPath, modelName)"
    """accuracy 
    contente = image base 64"""
    response["model"] = modelName
    response["state"] = "success"
    return JsonResponse(response)

@login_required
def models_dpgf_demo_ask(request):
    data = json.loads(request.body.decode('utf-8'))
    prompt = data.get('prompt', '') # question posée
    #model = DPGF(modelName="prix_prediction")
    prix = "model.ask(prompt)"
    response = {"content": prompt+" : " + str(prix),"state":"success"}
    return JsonResponse(response)


def chatapp_init_session(request):
    request.session['analyse_courante'] = ""
    request.session['fiches_selectionnees'] = []

def user_destination_dir(request):
    username = request.user.username if request.user.is_authenticated else 'visitor'
    destination_dir = os.path.join(get_absolute_path("sharepoint"), username)
    return destination_dir

@sync_to_async
def user_destination_dira(request):
    username = request.user.username if request.user.is_authenticated else 'visitor'
    destination_dir = os.path.join(get_absolute_path("sharepoint"), username)
    return destination_dir


def list_analyses(destination_dir,username):
    subdirectories = [d for d in os.listdir(destination_dir) if os.path.isdir(os.path.join(destination_dir, d)) and not d.startswith('_')]
    analyses_list = []
    if subdirectories:
        for i, caption in enumerate(subdirectories):
            analyses_dict = {"id": str(i + 1), "caption": caption, "type":"analyse", "parent_analyse":"", "memorized":True}
            analyses_list.append(analyses_dict)
        else:
            "pas de sous-répertoires"
    response_data={"username":username, "entries": analyses_list}
    return JsonResponse(response_data)


def list_analyse_files(destination_dir, analyse):
    files = os.listdir(destination_dir)
    file_list = []
    for i, caption in enumerate(files):
        # on évite les sous dossiers !
        file_path = os.path.join(destination_dir, caption)
        if os.path.isfile(file_path) and not (caption.endswith(".json") or caption.endswith(".txt")):
            file_dict = {"id": str(i + 1), "caption": caption, "type":"document", "parent_analyse":analyse}
            json_filename = caption + ".json"
            if os.path.isfile(os.path.join(destination_dir, json_filename)):
                file_dict["memorized"] = True
            else:
                # est-ce utile
                file_dict["memorized"] = True
            file_list.append(file_dict)
    response_data={"entries": file_list}
    return JsonResponse(response_data)



# REST JSON API

def chatapp_ajax_new_analyse(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            new_directory_name = data.get('analyse', '')
            user_dir = user_destination_dir(request)
            new_directory_path = os.path.join(user_dir, new_directory_name)

            if not new_directory_name:
                return JsonResponse({"error": "Le nom du répertoire ne peut pas être vide."}, status=400)

            if os.path.exists(new_directory_path):
                return JsonResponse({"error": "Le répertoire existe déjà."}, status=400)
            os.makedirs(new_directory_path)
            return list_analyses(user_dir,request.user.username) 
        except Exception as e:
            # Gérer les erreurs éventuelles ici
            return JsonResponse({"error": str(e)}, status=500)
    # Réponse en cas de méthode HTTP incorrecte
    return JsonResponse({"error": "Cette vue accepte uniquement les requêtes POST."}, status=405)


@login_required
def chatapp_upload(request):
    if request.method == 'POST' and request.FILES.getlist('files'):
        uploaded_files = request.FILES.getlist('files')
        analyse = request.POST.get('analyse')
        success_count = 0
        failure_count = 0
        for uploaded_file in uploaded_files:
            destination_dir = user_destination_dir(request)
            files_path = os.path.join(destination_dir, analyse)
            destination_path = os.path.join(files_path, uploaded_file.name)

            try:
                os.makedirs(destination_dir, exist_ok=True)
                with open(destination_path, 'wb+') as destination:
                    for chunk in uploaded_file.chunks():
                        destination.write(chunk)
                success_count += 1
            except Exception as e:
                failure_count += 1

        if success_count > 0:
            message = f'{success_count} fichiers téléchargés avec succès.'
        else:
            message = 'Échec du téléchargement des fichiers.'

        if failure_count > 0:
            message += f' {failure_count} fichiers n\'ont pas pu être téléchargés.'
            return JsonResponse({'message': 'Aucun fichier n\'a été téléchargé.'}, status=400)
        return list_analyse_files(files_path, analyse)

    return JsonResponse({'message': 'Aucun fichier n\'a été téléchargé.'}, status=400)




"""fonction pour supprimer un répertoire récursivement """
def delete_directory_recursive(repertoire):
    try:
        # Supprimer tous les fichiers à l'intérieur du répertoire
        for element in os.listdir(repertoire):
            element_path = os.path.join(repertoire, element)
            if os.path.isfile(element_path):
                os.remove(element_path)
            elif os.path.isdir(element_path): # Si c'est un sous-répertoire, récursivement supprimer son contenu
                delete_directory_recursive(element_path)

        # Supprimer le répertoire lui-même
        os.rmdir(repertoire)
        print(f"Le répertoire {repertoire} et son contenu ont été supprimés avec succès.")
    except Exception as e:
        print(f"Une erreur s'est produite : {e}")


"""vue pour supprimer un répertoire d'analyse"""
def chatapp_ajax_delete_analyse(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        analyse=data["analyse"]
        user_dir = user_destination_dir(request)
        directory_to_delete = os.path.join(user_dir, analyse)
        delete_directory_recursive(directory_to_delete)
        return list_analyses(user_dir,request.user.username) 



"""retieve analyse folders for current user"""
def chatapp_ajax_analyses(request):
    if request.method == 'GET':
        chatapp_init_session(request)
        return list_analyses(user_destination_dir(request),request.user.username)  




"""set current analisys and retieve files for current user/analyse"""
def chatapp_ajax_set_analyse(request):
    if request.method == 'POST':
        analyse = request.POST.get('analyse', None)
        request.session['analyse_courante'] = analyse
        request.session['fiches_selectionnees'] = []
        destination_dir = os.path.join(user_destination_dir(request), analyse)
        return list_analyse_files(destination_dir,analyse)


"""fusion des pdf dans le dossier d'analyse"""
def chatapp_ajax_fusion_pdf(request):
    if request.method == 'POST':
        analyse = request.POST.get('analyse', None)
        request.session['analyse_courante'] = analyse
        destination_dir = os.path.join(user_destination_dir(request), analyse)
        PdfUtils.folder_fusion(destination_dir,"fusion.pdf")
        return list_analyse_files(destination_dir,analyse)


"""supprime un pdf du dossier d'analyse"""
def chatapp_delete_file(request):
    if request.method == 'POST':
        data = json.loads(request.body.decode('utf-8'))

        
        analyse = data.get('analyse', '')
        fiches = data.get('entries', '')
        destination_dir = os.path.join(user_destination_dir(request), analyse)
        delete_file_path = os.path.join(destination_dir,fiches[0])
        deletefile(delete_file_path)
        deletefile(delete_file_path+".json")
        deletefile(delete_file_path+".json.txt")
        return list_analyse_files(destination_dir,analyse)


def deletefile(path):
    if os.path.exists(path):
        try:
            os.remove(path)
            return True
        except Exception as e:
            print(f"Une erreur s'est produite lors de la suppression du fichier : {e}")
            return False
    else:
        return False

"""
met à jour les informations de session : non utilisé jusqu'à présent
"""
def chatapp_set_fiches(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        request.session['analyse_courante'] = data["analyse"]
        request.session['fiches_selectionnees'] = data["analyse"]
        return JsonResponse(data)


"""if 'current_fiches' not in request.session:
        response_data = {"content": "choisir un fichier avec lequel discutter"}
        return JsonResponse(response_data)"""


def build_talk_response(message,state,json={}):
    response_data = {
        "content": f"""{message}""",
        "state":state,
        "json":json
        }
    return response_data







#############################################################################
""" CHROMA FUNCTIONS """
#############################################################################

def chroma_reset(request):
    
    import chromadb
    from chromadb.config import Settings
    path ="db/" 
    if os.path.isdir(path):
        client = chromadb.PersistentClient(path=path, settings=Settings(allow_reset=True))

        client.reset()  # Réinitialise la base de données
        state = True
    else : 
        state = False
    print()
    return JsonResponse({'status': 'success', 'message': f'Chroma reset : {state}'})










#############################################################################
"""CHAT LLM FUNCTIONS """
#############################################################################


def chatapp_llm(request):
    if request.method == 'POST':
        # Récupérer le nom du LLM depuis la requête AJAX
        data = json.loads(request.body.decode('utf-8'))
        llm_name = data.get('name', 'openai')
        
        if llm_name == "debug":
            request.session['llm_debug'] = True
        else : 
            request.session['llm_debug'] = False
        # Mettre à jour la session avec le modèle de langage sélectionné
        request.session['selected_llm'] = llm_name
        

        # Retourner une réponse JSON pour l'interface utilisateur
        return JsonResponse({'status': 'success', 'message': f'LLM sélectionné : {llm_name}'})
    
    return JsonResponse({'status': 'error', 'message': 'Méthode HTTP non supportée'}, status=405)


def chatapp_file_to_rag(request):
    if request.method == 'POST':
        # Récupérer le nom du LLM depuis la requête AJAX
        data = json.loads(request.body.decode('utf-8'))
        file_to_rag = data.get('name', '')
        
        if file_to_rag == "":
            request.session['file_to_rag'] = ""
        else : 
            request.session['file_to_rag'] = ".txt"
        # Mettre à jour la session avec le modèle de langage sélectionné
        request.session['file_to_rag'] = file_to_rag
        # Retourner une réponse JSON pour l'interface utilisateur
        return JsonResponse({'status': 'success', 'message': f'Utilisation du fichier {file_to_rag}'})

        
    

def gen_request(request):
    destination_dir = user_destination_dir(request)
    data = json.loads(request.body.decode('utf-8'))
    analyse = data.get('analyse', '') # nom du dossier d'analyse
    entries = data.get('entries', '') # liste des documents sélectionnés pour la question 
    prompt = data.get('prompt', '') # question posée
    history = data.get('history', '') # historique de conversation
    llm = request.session.get('selected_llm', 'openai')
    print(llm)
    if len(entries)<1:
        entries=["just.chat"]
    
    file_to_rag = request.session.get('file_to_rag', '')
    if file_to_rag =="":
        # path vers le pdf analysé
        pdf=os.path.join(destination_dir,analyse,entries[0])
    else:
        pdf=os.path.join(destination_dir,analyse,entries[0])
        if os.path.exists(pdf + file_to_rag):
            pdf = pdf + file_to_rag
    
    return analyse, entries, prompt, history, llm, pdf, data



def getjson_conversation(json_conversation_path):
    if os.path.exists(json_conversation_path):
        with open(json_conversation_path, "r", encoding="utf-8") as f:
            json_conversation = json.load(f)
    else:
        json_conversation={}
        
    return json_conversation


def chatapp_get_conversation(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    
    json_conversation_path = pdf+".json"
    response_data = getjson_conversation(json_conversation_path)
    # print(response_data)
    return JsonResponse(response_data)

def chatapp_save_conversation(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    
    json_conversation_path = pdf+".json"
    conversation_json = chat_save_conversation(src=json_conversation_path, history=history) 
    
    response_data = getjson_conversation(json_conversation_path)
    print(response_data)
    return JsonResponse(response_data)



# envoie une question sur le(s) documents sélectionnés (uniquement le premier en version demo, le mode multi-doc est à implémenter dans naldia)
# fait-on appel à des prompts préétablis ?

def chatapp_talk(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    
    if not history:
        history=[]
    # print(history)
    
    if request.session.get('llm_debug', False):
        # pdf = "llm_debug"
        # prompt = '\n\n'.join([lorem.paragraph() for _ in range(3)])
        pass
        
    if len(entries)<1:
        response_data = build_talk_response("Attention, veuillez choisir le(s) documents avec lesquels vous souhaitez discuter","danger")
        return JsonResponse(response_data)
    
    
    
    # s'agit-t-il d'un prompt préenregistré et quel est le format de réponse attendu (json / texte, ...)
    awaited_result = special_prompts(prompt) 
    
    if awaited_result["format"] == "text":
        response, conversation_json = chat(pdf=pdf, question = awaited_result["prompt"], history = history, llm=llm)
        response_data = build_talk_response(response,"warning", json=conversation_json)
    
    elif awaited_result["format"] == "json":
        try:
            history=[]
            response, conversation_json  = chat(pdf=pdf, question = awaited_result["prompt"], history = history, llm=llm)
            if prompt=="map_plot_pdf":
                jsonpath = pdf+".map.json"
                save_json(jsonpath, conversation_json)
            
            
            #donnees_json = json.loads(response) # Analyser la chaîne JSON en un objet Python
            return JsonResponse(response) # on renvoie la réponse JsonResponse avec les données JSON

        except json.JSONDecodeError as e:
            # La chaîne n'est pas au format JSON valide
            return JsonResponse({'error': 'La chaîne n\'est pas au format JSON valide'}, status=400)
    return JsonResponse(response_data)







"""
propose un sommaire d'analyse
"""
def chatapp_sommaire(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    
    if request.session.get('llm_debug', False):
        donnees_json = [
    {
        "title": "Objectif principal du document",
        "description": "Le document vise à fournir un Cahier des Clauses Techniques Particulières (CCTP) détaillé qui définit les spécifications techniques, les obligations des entreprises, les normes et règlements à suivre, ainsi que les prescriptions générales pour garantir une exécution conforme aux délais et exigences de qualité. Cela est crucial pour le bon déroulement des travaux tout en respectant les normes de santé."
    },
    {
        "title": "Public cible",
        "description": "Le public ciblé par ce document comprend principalement les entreprises de construction et d'ingénierie. Cela souligne l'importance de préparer un document technique qui répond à leurs besoins spécifiques pour soumissionner efficacement sur le lot n°4."
    },
    {
        "title": "Importance de la conformité aux normes",
        "description": "Le document souligne l'importance d'assurer que les travaux se conforment aux normes de santé et de sécurité en milieu hospitalier, ce qui est essentiel non seulement pour le bien-être des patients et du personnel mais aussi pour éviter des complications légales éventuelles."
    },
    {
        "title": "Continuité des services",
        "description": "Le CCTP prend en compte la nécessité de maintenir la continuité des services existants durant la phase d'extension, ce qui est vital dans un établissement de santé pour assurer que les soins aux patients ne soient pas interrompus."
    },
    {
        "title": "Collaboration avec divers acteurs",
        "description": "Le document implique plusieurs parties prenantes, notamment le Maître d’Ouvrage (Centre Hospitalier de Béziers) et le Bureau d’Études (EREN Ingénierie), ce qui nécessite une coordination efficace pour le succès du projet."
    },
    {
        "title": "Prescriptions générales et spécifiques",
        "description": "Le CCTP inclut des prescriptions générales et spécifiques concernant les travaux de plomberie, Chauffage, Ventilation, et Climatisation, ce qui permet d'établir des attentes claires pour les entreprises soumissionnaires."
    }
]
        return JsonResponse(donnees_json, safe=False)
    
    response = extract_insights(pdf=pdf, pages = prompt, history = history, llm=llm)
    return JsonResponse(response, safe=False)



def chatapp_delete_conversation_entry(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    
    delete_entry(pdf, prompt)
    response_data = build_talk_response(f"Entrée effacée de l'historique : {prompt}","danger")
    return JsonResponse(response_data)
    




def chatapp_enhance(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    originalText = data.get('originalText', '') # texte a augmenter
    originalQuestion = data.get('originalQuestion', '') # question initiale
    llm_debug = request.session.get('llm_debug', False)
    
    
    
    if llm_debug:
        pass
         
    if len(entries)<1:
        response_data = build_talk_response("Attention, veuillez choisir le(s) documents avec lesquels vous souhaitez discuter","danger")
        return JsonResponse(response_data)
    
    response, conversation_json = chat_enhance(originalQuestion = originalQuestion, originalText= originalText, pdf=pdf, question = prompt, history = history, llm=llm)
    response_data = build_talk_response(response,"warning", json=conversation_json)
    return JsonResponse(response_data)



from docx import Document
from bs4 import BeautifulSoup
def ajouter_contenu_html(doc, contenu_html):
    # Utilise BeautifulSoup pour parser le contenu HTML
    soup = BeautifulSoup(contenu_html, 'html.parser')

    for element in soup.descendants:
        if element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'strong':
            run = doc.add_paragraph().add_run(element.text)
            run.bold = True
        elif element.name == 'em':
            run = doc.add_paragraph().add_run(element.text)
            run.italic = True
        elif element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListNumber')
        # Vous pouvez ajouter d'autres balises si nécessaire
        
def generer_document_word(donnees, nom_fichier="document_genere.docx"):
    # Crée un nouveau document Word
    doc = Document()
    conversation = []
    # Boucle à travers chaque entrée du tableau pour ajouter des chapitres
    for entree in donnees:
        titre, contenu = entree  # Décompose l'entrée en titre et contenu
        # Ajoute le titre comme un en-tête de niveau 1
        conversation.append({"title":titre,"description": contenu})
        doc.add_heading(titre, level=1)
        # Ajoute le contenu comme un paragraphe
        ajouter_contenu_html(doc, contenu)
    # Enregistre le document avec le nom spécifié
    
    conversation_json  = {
        "conversationPath":nom_fichier+".json",
        "conversation": conversation
        
    }
    doc.save(nom_fichier)
    print(f"Document généré avec succès : {nom_fichier}")
    
    
    # Enregistre l'objet JSON dans un fichier
    nom_fichier_json = nom_fichier + ".json"
    save_json(nom_fichier_json, conversation_json)
    print(f"Conversation sauvegardée en JSON : {nom_fichier_json}")


def save_json(jsonpath, jsontext):
     with open(jsonpath, "w", encoding="utf-8") as f:
        json.dump(jsontext, f, indent=4, ensure_ascii=False)

    
def chatapp_word(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    
    if len(entries)<1:
        response_data = build_talk_response("Attention, veuillez choisir le(s) documents avec lesquels vous souhaitez discuter","danger")
        return JsonResponse(response_data)
    
    
    filepath=os.path.join(user_destination_dir(request),analyse,f"{prompt}.docx")
    generer_document_word(history, filepath)
    response_data = build_talk_response("Le word a été généré","info")
    return JsonResponse(response_data)




# todo => intégrer les prompts préétablis
# extraction GPS, Diagrammes circulaire, extraction sommaires et informations chiffrées ...
def special_prompts(prompt):
    if prompt.startswith("map_plot_"): 
        #extraire des informations de géolocalisation d'un document
        expectedformat="json"
    else:
        expectedformat="text"
    
    awaited_result = {"prompt":prompt,"format":expectedformat}
    return awaited_result


# todo => améliorer la présentation des réponses OpenAI
def format_text(text):
    return markdown.markdown(text)





# mémorize document (mode mulit-doc non implémenté)
def chatapp_summarize(request):
    analyse, entries, prompt, history, llm, pdf, data = gen_request(request=request)
    
    llm_debug = request.session.get('llm_debug', False)
    
    if llm_debug: 
        prompt = '\n\n'.join([lorem.paragraph() for _ in range(3)])
        response_data = build_talk_response(f"chatapp_summarize : nombre de pages : {prompt}","warning")
        return JsonResponse(response_data)

    if len(entries)<1:
        response_data = build_talk_response("Attention, veuillez choisir le(s) documents avec lesquels vous souhaitez discuter","danger")
        return JsonResponse(response_data)
    
    response = chat_summarize(pdf=pdf, pages = prompt, history = history, llm=llm)
    response_data = build_talk_response(response,"warning")
    return JsonResponse(response_data)




# retour le cpntenu du ficher de vectorisation pour l'afficher dans le navigateur
def afficher_resume_vectorisation(request, analyse, nom_fichier):
    file_path = os.path.join(settings.BASE_DIR, 'chatapp','sharepoint', request.user.username, analyse, nom_fichier+".txt")
    
    json_conversation_path = os.path.join(settings.BASE_DIR, 'chatapp','sharepoint', request.user.username, analyse, nom_fichier+".json")
    json_conversation = getjson_conversation(json_conversation_path)
    
    
    # Vérifiez que le fichier existe avant de le renvoyer
    # printc(file_path,bcolors.FAIL)

    if os.path.exists(file_path):
        # Ouvrez le fichier PDF et renvoyez-le en tant que réponse HTTP
        with open(file_path, 'r', encoding='utf-8') as fichier:
            # Lire le contenu du fichier
            response_data = fichier.read()
        response_data = build_talk_response(format_text(response_data),"success", json=json_conversation)
        return JsonResponse(response_data)
    else:
        response_data = build_talk_response("le résumé du fichier n'existe pas, lancez la procédure","warning", json=json_conversation)
        return JsonResponse(response_data)





# non utilisé, permet de renvoyer le fichier sélectionner
def afficher_ressources(request, analyse, nom_fichier):
    file_path = os.path.join(settings.BASE_DIR, 'chatapp','sharepoint', request.user.username, analyse, nom_fichier)
    # Vérifiez que le fichier existe avant de le renvoyer
 
    if os.path.exists(file_path):
        # Ouvrez le fichier PDF et renvoyez-le en tant que réponse HTTP
        with open(file_path, 'rb') as f:
            response = FileResponse(f, as_attachment=True)
            return response
    else:
        # Gérez le cas où le fichier n'existe pas, par exemple, en renvoyant une erreur 404.
        from django.http import Http404
        raise Http404("Le fichier PDF n'existe pas.")